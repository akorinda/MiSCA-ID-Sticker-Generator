"""Create scannable code images

Packages: sys, os, pandas, qrcode, PIL, datetime, tkinter

"""

__authors__ = "Andrew Korinda"
__copyright__ = "Copyright 2021, Midland Mountain Bike Crew"
__credits__ = ["Andrew Korinda"]
__license__ = "GPL-3.0-or-later"
__version__ = "0.1"
__maintainer__ = "https://github.com/akorinda/MiSCA-ID-Sticker-Generator"
__status__ = "Prototype"


# Imports
import sys
import subprocess, os, platform
import time

import pandas
import qrcode
import threading, queue
from PIL import Image, ImageDraw, ImageFont
from tkinter import filedialog
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_TAB_ALIGNMENT, WD_LINE_SPACING


# Constants

def qr_create(txt_data):
    qr = qrcode.QRCode(
        version=4,
        error_correction=qrcode.constants.ERROR_CORRECT_H,
        box_size=40,
        border=7
    )

    qr.add_data(txt_data)
    qr.make(fit=True)

    qr_img = qr.make_image(
        fill='black',
        back_color='white'
    )

    return qr_img


def frame_create(group):
    frame_img = Image.new(
        'RGB',
        (2100, 2611),
        color=group
    )

    frame_img = frame_img.resize((2000, 2511))

    border_img = Image.new(
        'RGB',
        (2100, 2611),
        color='white'
    )

    border_img.paste(
        frame_img,
        (50,
         50),
        mask=None
    )

    return border_img


def set_font(default_font_size):
    default_font = ImageFont.truetype(
        font='arialbd.ttf',
        size=default_font_size
    )

    return default_font


def handle_logo(int_path, base_img):
    # valid_selection:bool = False
    # print('\nWith QR codes you can add a logo to the center.')
    # print('A downside is the qr code will have less error handling.')
    
    # while not valid_selection:
    #     confirm_desire = input('Will you like to add a logo? [Y,N]: ')
        
    #     if confirm_desire.lower() in ['y', 'yes']:
    #         logo_file = filedialog.askopenfilename(
    #             title='Open Center Logo',
    #             filetypes=[("Image file", "*.jpg *.jpeg *.png")],
    #             initialdir=int_path
    #         )
    #         valid_selection = True
    #     elif confirm_desire.lower() in ['n', 'no']:
    #         logo_file = ''
    #         valid_selection = True

    # try:
    #     logo_img = Image.open(logo_file, mode='r')
    #     logo_dim = round((0.07 * base_img.size[0] * base_img.size[1]) ** 0.5)
    #     logo_img = logo_img.resize((logo_dim, logo_dim))
    # except AttributeError:
    logo_n = 1
    logo_img = Image.new('RGB', (logo_n, logo_n))

    return logo_img


def display_options(prime_list, extend_list, x_pos, options_size):
    input_index = 1
    options_list = []
        
    for x in range(x_pos, min(x_pos + options_size, len(prime_list))):
        print(f"     {input_index}. {prime_list[x]}")
        input_index += 1
        options_list.append(prime_list[x])
        
    for x in range(len(extend_list)):
        print(f"{extend_list[x]}")
        
    return options_list


def selection_evaluation(prime_inputs, extend_inputs, user_input,
                         output_list,
                         list_step:int, list_max:int, x_pos:int,
                         min_selections:int):
    global universal_input_options
    valid_input:bool = True
    continue_requests:bool = True
    user_input = user_input.lower()
    
    if user_input in universal_input_options:
        print('\n\nIn the future this will goto universal evaluation function but for now:')
        print('\n\nThe user choose to quit the program')
        sys.exit()
    elif user_input in extend_inputs:
        if user_input in ["more", '0']:
            x_pos += list_step
            if x_pos >  list_max: x_pos = 0
            valid_input = False
        elif user_input in ['done', '99']:
            if len(output_list) >= min_selections:
                continue_requests = False
            else:
                print(f'\nAt least {min_selections} selection(s) must be made, currently there are {len(output_list)}.')
                continue_requests = True
                valid_input = False
        elif user_input in ['clear list', '*']:
            output_list = []
            x_pos = 0
        elif user_input in ['backspace', '~']:
            if len(output_list) > 0:
                output_list.pop()
    elif user_input in prime_inputs[0]: # Index given
        output_list = add_to_info_list(output_list, 
                                       prime_inputs[1], 
                                       x_pos + int(user_input) - 1)
    elif user_input in prime_inputs[1]: # String given
        user_input = list(prime_inputs[1]).index(user_input)
        output_list = add_to_info_list(output_list, 
                                       prime_inputs[1],
                                       user_input)
    else:
        valid_input = False
        print("\n!!!!! Selection was not recognized !!!!!")
        print("Options available at this stage are:")
        print("  Number without period or neighboring text")
        for x in range(0, len(extend_inputs), 2):
            try:
                print(f'  "{extend_inputs[x]}" or "{extend_inputs[x+1]}"')
            except IndexError:
                print(f'  "{extend_inputs[x]}"')
        print('  Type "h" or "help" at input for more help information')
                
    return output_list, x_pos, continue_requests, valid_input


my_queue = queue.Queue()

def storeInQueue(f):
    def wrapper(*args):
        global my_queue
        my_queue.put(f(*args))
    return wrapper


@storeInQueue
def load_excel(declared_file):
    try:
        book = pandas.ExcelFile(declared_file)
    except PermissionError:
        print(os.path.basename(declared_file) + ' is not accessible. Likely in use by another application.')
        book = 'PermissionError'
    except (AssertionError, FileNotFoundError):
        print('No rider list was selected. Exiting')
        book = 'AssertionError'
    return book


def spinner(text):
    if len(text) < 17:
        text = text + '.'
    else:
        text = 'Loading...'
    print(f'\r{text}', end='', flush=True)
    return text


def data_file(def_title):
    global universal_input_options, my_queue
    
    selected_file = filedialog.askopenfilename(
        title=def_title,
        filetypes=[("Excel file", "*.xlsx *.xls")]
    )
    
    if len(selected_file) > 0:
        load_indicator = "Loading.."
    
        t = threading.Thread(target=load_excel, args=(selected_file, ))
        t.start()
        while t.is_alive():
            load_indicator = spinner(load_indicator)
            time.sleep(0.1)
    
        data_book = my_queue.get()
    else:
        print('No rider list was selected. Exiting')
        sys.exit(0)
        
    if data_book == 'PermissionError':
        sys.exit(1)
    elif data_book == 'AssertionError':
        sys.exit(0)
    elif isinstance(data_book, pandas.ExcelFile):
        print('[Complete]')
    else:
        print('\n!!!!! Unrecognized data_book Returned !!!!!')
        print(f'{data_book}')
        sys.exit(1)

    # Select the worksheet from the workbook
    sheet_select = []
    iList_size = 9
    xStart = 0
    continue_selection: bool = True

    while continue_selection:
        if len(data_book.sheet_names) > 1:
            print("\nWhich worksheet contains the rider data?")
            xtraList = ["     0. Show more columns [More]"]

            dList = display_options(data_book.sheet_names,
                                    xtraList,
                                    xStart,
                                    iList_size)

            info_selection = input(f"Choice (1:{min(len(dList), iList_size)}): ")
            if info_selection.lower() in universal_input_options:
                print('\n\nThe user choose to quit the program')
                sys.exit()
            elif info_selection.lower() == "more" or info_selection == '0':
                xStart += iList_size
                if xStart > len(data_book.sheet_names): xStart = 0
            elif info_selection.lower() in map(lambda x: x.lower(), data_book.sheet_names):
                info_selection = [var.lower() for var in dList].index(info_selection)
                sheet_select = add_to_info_list(sheet_select,
                                                data_book.sheet_names,
                                                info_selection)
                continue_selection = False
            elif info_selection in map(str, range(1, iList_size + 1)):
                sheet_select = add_to_info_list(sheet_select,
                                                dList,
                                                int(info_selection) - 1)
                continue_selection = False
            else:
                print("\n!!!!! Selection was not recognized !!!!!")
                print("Options available at this stage are:")
                print("Number without period or column name")
                print("0 or More")
        else:
            sheet_select = data_book.sheet_names
            continue_selection = False

    try:
        selected_sheet = sheet_select[0]
    except KeyError:
        print("\n!!!!! A Selection was not found !!!!!")
        print(f"Received: {sheet_select}")

    data_rows = data_book.parse(sheet_name=selected_sheet)

    return selected_file, selected_sheet, data_rows


def code_document_request(int_path):
    try:
        doc_name = filedialog.asksaveasfilename(
            title='QR Document Save Location',
            initialdir=int_path,
            defaultextension='*.docx',
            filetypes=[('Word file', '*.docx')],
            confirmoverwrite=True
        )

        if doc_name == '':
            print('No output path was selected. Exiting')
            sys.exit(0)
    except FileNotFoundError:
        print('The system cannot find the path specified. Exiting')
        sys.exit(0)

    doc_path = os.path.dirname(doc_name)
    if not os.path.exists(doc_path):
        try:
            os.makedirs(doc_path)
        except PermissionError:
            print(doc_path + ' is not accessible.')
            sys.exit(1)
        except AssertionError:
            print('No output path was selected. Exiting')
            sys.exit(0)
        except FileNotFoundError:
            print('The system cannot find the path specified. Exiting')
            sys.exit(0)

    return doc_name


def add_to_info_list(current_list, available_list, add_index:int):
    add_item = list(available_list)[add_index]
    current_list.append(add_item)
        
    return current_list


def image_generation(info, text_top, text_bottom):
    group = 'white'
    text = text_top + '\n' + text_bottom
    
    # Create layers
    layer_qr = qr_create(info)
    layer_frame = frame_create(group)
    frame_w, frame_h = layer_frame.size
    
    # Resize QR and logo per border definition
    layer_qr = qr_create(group)
    # Option for adding a logo to the center of the qr code
    layer_logo = handle_logo(os.path.dirname(rider_file), layer_qr)
    # % Dimensions
    logo_w, logo_h = layer_logo.size
    qr_w, qr_h = layer_qr.size
    qr_buffer = round((frame_w - qr_w) / 2)
    
    # Stack layers
    rider_image = layer_frame
    rider_image.paste(
        layer_qr,
        (qr_buffer, qr_buffer),
        mask=None
    )
    rider_image.paste(
        layer_logo,
        (round((frame_w - logo_w) / 2),
         round((qr_h - logo_h) / 2) + qr_buffer),
        mask=None
    # Add rider name text
    )
    start_size = 240
    font_size = start_size
    font = set_font(font_size)

    if len(text_top) > len(text_bottom):
        name_limit = text_top
    else:
        name_limit = text_bottom

    # Make sure the name fits
    while True:
        font_size = font_size - 20
        font = set_font(font_size)
        text_vert_pos = frame_h / 2 + qr_h / 2 - qr_buffer / 2  # was + qr_buffer

        if (font.getsize(name_limit)[0] < qr_w) or (font.size <= 0):
            break

    name_color = 'black'

    rider_label = ImageDraw.Draw(rider_image)
    rider_label.text(
        (round(frame_w / 2),
         text_vert_pos),
        text,
        fill=name_color,
        anchor='mm',
        align='center',
        font=font
    )

    # Outputs
    rider_image = rider_image.resize((300, 373))

    return rider_image


# Following the template for Avery Presta 94103, 1"x1" sticker labels
def qr_line_format(paragraph_to_format):
    paragraph_format = paragraph_to_format.paragraph_format
    tab_stops = paragraph_format.tab_stops
    for x in range(6):
        tab_stops.add_tab_stop(Inches(0.5 + 1.23*x), WD_TAB_ALIGNMENT.CENTER)

    paragraph_format.space_before = Inches(0)
    paragraph_format.space_after = Inches(0.25)
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    return paragraph_to_format


# Following the template for Avery Presta 94103, 1"x1" sticker labels
def qr_doc_format(document_to_format):
    section = document_to_format.sections[0]
    section.left_margin = Inches(0.71)
    section.right_margin = Inches(0.31)
    section.top_margin = Inches(0.63)
    section.bottom_margin = Inches(0.57)
    section.gutter = Inches(0)

    return document_to_format


def exception_file_open(requested_document):
    print(requested_document + ' is not accessible. Likely in use by another application.')
    sys.exit(1)


# Run the program
if __name__ == '__main__':
    divider = '--------------------------------------------------------------------------'
    clear_command = "cls" if platform.system() == "Windows" else "clear"
    os.system(clear_command)

    print("\n----------- Welcome to Making QR Code Rider Sheets -----------")
    print("- 'X' to return not implemented")
    print("- 'H' for help text not implemented")
    print("- 'Q' to quit the program not implemented")
    universal_input_options = ['x', 'exit',
                               'h', 'help',
                               'q', 'quit']
    print("\nThe order of operations is not yet configurable.\nIt is:")
    print("1) Select the rider data file")
    print("2) Select worksheet if multiple are in the file")
    print("3) Select the columns which will be part of the qr code or barcode")
    print("4) Filter options are presented")
    print("5) ***Selection of qr code or barcode not implemented***")
    print("6) Create or Select an output Word document")
    print("7) Images are created and placed in the Word document")
    input('\nPress Enter to continue...')
    print(divider)
    
    # ToDo: resolved error in PIL Image when from tkinter import * and Tk() is used
    # root = Tk()  # pointing root to Tk() to use it as Tk() in program.
    # root.withdraw()  # Hides small tkinter window.
    # root.attributes('-topmost', True)  # Opened windows will be active. above all windows despite of selection.

    # Ask for xlsx rider list
    print("\n\nProvide the file with the rider list. An Excel formant, *.xls or *.xlsx is expected")
    rider_file, rider_sheet, riders = data_file('Open Rider List')
    print(f'Using "{os.path.basename(rider_file)}",\n  Sheet "{rider_sheet}" for rider information')
    # Drop columns without any data
    riders = riders.dropna(axis='columns',
                           how='all')

    print(divider)
    print("\n\nWhat information should be in the QR code, in order?")
    rider_info = []
    iList_size = 9
    xStart = 0
    continue_selection:bool = True
    delim = '\t'
    while continue_selection:
        print(f"\nCurrent information: {delim.join(rider_info)}")
        print("Next item to include:")
        xtraList = ["     0. Show more columns [More]",
                    "     99. Done with selections [Done]",
                    "     ~. Remove last selection [Backspace]",
                    "     *. [Clear list]"]
        
        dList = display_options(riders.columns,
                                xtraList,
                                xStart,
                                iList_size)
        
        info_selection = input(f"Choice (1:{min(len(dList), iList_size)}): ")
        print('\n')
        if info_selection.lower() in universal_input_options:
            print('\n\nThe user choose to quit the program')
            sys.exit()
        elif info_selection.lower() == "more" or info_selection == '0':
            xStart += iList_size
            if xStart > riders.columns.size: xStart = 0
        elif info_selection.lower() == 'done' or info_selection == '99':
            if len(rider_info) > 0:
                continue_selection = False
            else:
                print('\nAt least one selection must be made, currently there are none.')
        elif info_selection.lower() == 'clear list' or info_selection == '*':
            rider_info = []
            xStart = 0
        elif info_selection.lower() == 'backspace' or info_selection =='~':
            if len(rider_info) > 0:
                rider_info.pop()
        elif info_selection.lower() in map(lambda x:x.lower(), riders.columns):
            info_selection = [var.lower() for var in riders.columns].index(info_selection)
            rider_info = add_to_info_list(rider_info, 
                                          riders.columns,
                                          info_selection)
        elif info_selection in map(str, range(1,iList_size + 1)):
            rider_info = add_to_info_list(rider_info, 
                                          dList, 
                                          int(info_selection) - 1)
        else:
            print("\n!!!!! Selection was not recognized !!!!!")
            print("Options available at this stage are:")
            print("Number without period or column name")
            print("0 or More")
            print("99 or Done")
            print("~ or Backspace")
            print("* or Clear list")
    
    # Remove rows without data in key column
    riders = riders.dropna(subset=rider_info)
    # Drop all rows where the key column is the only coulmn with data
    riders = riders.dropna(thresh=2)

    print(divider)
    print('\n\nWhat if any filters should be applied?')
    filter_dict = {}
    iList_size = 9
    continue_selection:bool = True
    valid_filter:bool = False
    delim = '\n  '
    delim_filters = ': '
    delim_select = ', '
    xtraList = [["     0. Show more columns [More]",
                 "     99. Done with selections [Done]",
                 "     ~. Remove last selection [Backspace]",
                 "     *. [Clear list]"],
                ['0', 'more',
                 '99', 'done',
                 '~', 'backspace',
                 '*', 'clear list']]
    while continue_selection and not valid_filter:
        filter_list = []
        xStart = 0
        # Select a column to filter
        while not valid_filter:
            filter_info = [list(filter_dict.keys()), list(filter_dict.values())]
            for x in range(len(filter_info[0])):
                try:
                    filter_list.append(delim_filters.join([filter_info[0][x],
                                                           delim_select.join(filter_info[1][x])]))
                except IndexError:
                    filter_list = []
                except NameError:
                    filter_list = []
                    try:
                        filter_list.append(delim_filters.join([filter_info[0][x],
                                                               delim_select.join(filter_info[1][x])]))
                    except IndexError:
                        filter_list = []
                    
            if len(filter_list) > 0:
                print(f"\nCurrent filters ↓↓↓\n  {delim.join(filter_list)}")
                print("Next column to include:")
            else:
                print('Column to filter:')
             
            dList = display_options(riders.columns,
                                    xtraList[0],
                                    xStart,
                                    iList_size)
            
            info_selection = input(f"Choice (1:{min(len(dList), iList_size)}): ")
            
            info_ints = list(map(str, range(1,iList_size + 1)))
            info_names = list(map(lambda x:x.lower(), riders.columns))
            
            filter_info[0], xStart, continue_selection, valid_filter = selection_evaluation([info_ints,info_names],
                                                                                            xtraList[1],
                                                                                            info_selection,
                                                                                            filter_info[0],
                                                                                            iList_size,
                                                                                            riders.columns.size,
                                                                                            xStart,
                                                                                            0)
        if valid_filter and continue_selection:
            filter_append = riders.columns[list(info_names).index(filter_info[0][-1])]
        else:
            break
        try:
            filter_dict[filter_append]
        except KeyError:
            filter_dict[filter_append] = []
        # filter_info[0][-1] = riders.columns[list(info_names).index(filter_info[0][-1])]
        # Now that the column is valid, select a valid value within that column
        valid_filter = False
        xStart = 0
        while not valid_filter or continue_selection:
            valid_filter = False
            print(f'\nSelect a value to filter for {filter_append}:')
            
            try: 
                filter_elements = list(map(lambda x:str(x),
                                           list(riders[filter_append].drop_duplicates())))
            except KeyError:
                print(f'"{filter_info[0][-1]}" was not found in {riders.columns}')
                sys.exit()
            
            dList = display_options(filter_elements,
                                    xtraList[0],
                                    xStart,
                                    iList_size)
            
            info_selection = input(f"Choice (1:{min(len(dList), iList_size)}): ")
            
            info_ints = list(map(str, range(1,iList_size + 1)))
            info_names = list(map(lambda x:x.lower(), filter_elements))
                        
            filter_info[1], xStart, continue_selection, valid_filter = selection_evaluation([info_ints,info_names],
                                                                                            xtraList[1],
                                                                                            info_selection,
                                                                                            filter_info[1],
                                                                                            iList_size,
                                                                                            len(filter_elements),
                                                                                            xStart,
                                                                                            1)
            
            if valid_filter and continue_selection:
                selection_append = filter_elements[list(info_names).index(filter_info[1][-1])]
                filter_dict[filter_append].append(selection_append)
                # filter_info[1][-1] = filter_elements[list(info_names).index(filter_info[1][-1])]
        continue_selection, valid_filter = True, False
                
    # Apply the filter
    # Brute force method
    try:
        filter_info = [list(filter_dict.keys()), list(filter_dict.values())]
        query = ''
        delim_or = ' or '
        delim_and = ' and '
        delim_eval = ' == '
        for x in range(len(filter_info[0])):
            if x == 0:
                query = "'"
            else:
                query = query + delim_and
            for y in range(len(filter_info[1][x])):
                if y == 0:
                    delim_or = '('
                else:
                    delim_or = ' or '
                query = query + delim_or + '`' + filter_info[0][x] + '`' + delim_eval + '"' + filter_info[1][x][y] + '"'
                if y == len(filter_info[1][x]) - 1:
                    query = query + ')'
        query = query + "'"
        riders = riders.query(eval(query))
        riders = riders.reset_index(drop=True)
    except SyntaxError:
        print('No filter was recognized, continuing')
        riders = riders # Apply no filter

    # ToDo: Check there is data remaining, if not return to filter selection

    # Ask for text rows
    print(divider)
    selected_lines = ['','']
    place_name = ['top', 'bottom']
    xtraList = [["     0. Show more columns [More]",
                 "     99. No text [Done]"],
                 ['0', 'more',
                  '99', 'done']]
    for x in range(2):
        xStart = 0
        iList_size = 9
        continue_selection = True
        valid_column = False
        while continue_selection and not valid_column:
            print(f"\nSelect the {place_name[x]} line of text")
            print("Column to use:")
            dList = display_options(riders.columns,
                                    xtraList[0],
                                    xStart,
                                    iList_size)
            
            info_selection = input(f"Choice (1:{min(len(dList), iList_size)}): ")
            
            info_ints = list(map(str, range(1,iList_size + 1)))
            info_names = list(map(lambda x:x.lower(), riders.columns))
                
            text_select, xStart, continue_selection, valid_column = selection_evaluation([info_ints,info_names],
                                                                                         xtraList[1],
                                                                                         info_selection,
                                                                                         list(selected_lines[x]),
                                                                                         iList_size,
                                                                                         riders.columns.size,
                                                                                         xStart,
                                                                                         0)
            time.sleep(0.01)
        if valid_column:
            selected_lines[x] = riders.columns[list(info_names).index(text_select[0])]    
        else:
            selected_lines[x] = ''

    # Select where final documents gets saved
    print(divider)
    print("\nCreate a document name for the output codes.")
    code_doc = code_document_request(os.path.dirname(os.path.realpath(__file__)))
    print(f"Using {code_doc} as the output document")
    code_path = os.path.dirname(code_doc)

    # How many copies of each code should be put in the document?
    correct_format = False
    while not correct_format:
        try:
            code_qty = input('\nHow many copies should be made for each participant?: ')
            if len(code_qty) > 0:
                code_qty = int(code_qty)
                correct_format = True
                if code_qty <= 0:
                    print('\nNo codes printed based on requested quantity.')
                    sys.exit(0)
        except ValueError:
            print("\nEnter an integer value.")
    print('\nCreating document...')

    try:
        document = Document()
        document.add_paragraph()
        document.add_section()
        # document.save(output_document)
        # document = Document(output_document)
    except PermissionError:
        exception_file_open(code_doc)

    document = qr_doc_format(document)
    paragraph = document.paragraphs[0]
    paragraph = qr_line_format(paragraph)

    qr_count = 0
    paragraph.add_run('\t')

    # ToDo: make range of riders selectable
    # ToDo: make the code type selectable
    # ToDo: handle people with the same name

    delim = '\t'
    riders_codes = riders[riders.columns.intersection(rider_info)].values.tolist()
    rider_dict = {}
    for idx_rider in range(len(riders)):
        rider_code = delim.join(str(var) for var in riders_codes[idx_rider])
        try:
            first_text = riders[selected_lines[0]][idx_rider].title()
        except KeyError:
            first_text = ''
        try:
            second_text = riders[selected_lines[1]][idx_rider].title()
        except KeyError:
            second_text = ''
        rider_text = second_text + ' ' + first_text
        rider_dict.update({rider_text: [rider_code, first_text, second_text]})

    idx_rider = 0
    for key_rider in sorted (rider_dict.keys()):
        rider_file = os.path.join(code_path, key_rider + '.png')
        
        rider_qr = image_generation(rider_dict.get(key_rider)[0],
                                    rider_dict.get(key_rider)[1],
                                    rider_dict.get(key_rider)[2])
        rider_qr.save(rider_file)

        # ToDo: show a progress meter
        for copy_number in range(1, code_qty + 1):
            run = paragraph.add_run()
            run.add_picture(rider_file, height=Inches(1.0))
            if idx_rider+1 % 6 == 0:
                paragraph = document.add_paragraph()
                paragraph = qr_line_format(paragraph)
                paragraph.add_run('\t')
            else:
                paragraph.add_run('\t')

        try:
            os.remove(rider_file)
        except OSError:
            pass

    try:
        # Save what has been created as specified
        document.save(code_doc)

        # Open the created document
        if platform.system() == 'Darwin':  # macOS
            subprocess.call(('open', code_doc))
        elif platform.system() == 'Windows':  # Windows
            os.startfile(code_doc)
        else:  # linux variants
            subprocess.call(('xdg-open', code_doc))
    except PermissionError:
        exception_file_open(code_doc)

    input('Process complete. Press enter to close the program or close the window as normal.')
    sys.exit()