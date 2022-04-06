"""Create printable document from scannable code images

Packages used: sys, os, docx, datetime

"""

# Metadata
__license__ = "GPL-3.0-or-later"
__maintainer__ = "https://github.com/akorinda/MiSCA-ID-Sticker-Generator"

# Imports
import os
import sys
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_TAB_ALIGNMENT, WD_LINE_SPACING
from datetime import date


# Following the template for Avery Presta 94103, 1"x1" sticker labels
def qr_line_format(paragraph_to_format):
    paragraph_format = paragraph_to_format.paragraph_format
    tab_stops = paragraph_format.tab_stops
    for x in range(6):
        tab_stops.add_tab_stop(Inches(0.5 + 1.23*x), WD_TAB_ALIGNMENT.CENTER)

    paragraph_format.space_before = Inches(0)
    paragraph_format.space_after = Inches(0.25)
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    return paragraph_to_format


# Following the template for Avery Presta 94103, 1"x1" sticker labels
def qr_doc_format(document_to_format):
    section = document_to_format.sections[0]
    section.left_margin = Inches(0.71)
    section.right_margin = Inches(0.31)
    section.top_margin = Inches(0.63)
    section.bottom_margin = Inches(0.57)
    section.gutter = Inches(0)

    return document_to_format


def code_doc_from_dict(code_dict, output_document, *, qty_copies=1, input_document=''):
    if len(input_document) > 0:
        try:
            document = Document(input_document)
        except PermissionError:
            exception_file_open(input_document)
    else:
        try:
            document = Document()
            document.add_paragraph()
            document.add_section()
            #document.save(output_document)
            #document = Document(output_document)
        except PermissionError:
            exception_file_open(output_document)

    document = qr_doc_format(document)
    paragraph = document.paragraphs[0]
    paragraph = qr_line_format(paragraph)

    qr_count = 0
    paragraph.add_run('\t')
    for key_i in sorted (code_dict.keys()):
        for x in range(1, qty_copies + 1):
            run = paragraph.add_run()
            run.add_picture(code_dict.get(key_i), height=Inches(1.0))
            qr_count += 1
            if qr_count % 6 == 0:
                paragraph = document.add_paragraph()
                paragraph = qr_line_format(paragraph)
                paragraph.add_run('\t')
            else:
                paragraph.add_run('\t')

    try:
        document.save(output_document)
    except PermissionError:
        exception_file_open(output_document)


# Main function of this script to compile all png files in the given directory
def qr_doc_create(input_document, output_document, image_path, qty_copies):
    try:
        document = Document(input_document)
    except PermissionError:
        exception_file_open(input_document)
    document = qr_doc_format(document)

    paragraph = document.paragraphs[0]
    paragraph = qr_line_format(paragraph)

    qr_path = image_path
    qr_images = []

    for file in os.listdir(qr_path):
        if file.endswith('.png'):
            qr_images.append(file)

    qr_count = 0
    paragraph.add_run('\t')
    for QR_Code in qr_images:
        for x in range(1,qty_copies + 1):
            run = paragraph.add_run()
            run.add_picture(qr_path + '/' + QR_Code, height=Inches(1.0))
            qr_count += 1
            if qr_count % 6 == 0:
                paragraph = document.add_paragraph()
                paragraph = qr_line_format(paragraph)
                paragraph.add_run('\t')
            else:
                paragraph.add_run('\t')

    try:
        document.save(output_document)
    except PermissionError:
        exception_file_open(output_document)


# Error handling when a file is in use by another program
def exception_file_open(requested_document):
    print(requested_document + ' is not accessible. Likely in use by another application.')
    sys.exit(1)


if __name__ == '__main__':
    input_file = 'Python Gen QR Printout.docx'
    output_file = 'QR-Full_' + date.today().strftime('%Y-%m-%d') + '.docx'
    image_pngs = os.path.dirname(__file__) + '/QR-Export/2021-08-18'
    copies = 3
    qr_doc_create(input_file, output_file, image_pngs, copies)
    print(__name__ + ' has compiled the images into ' + output_file + '.')
